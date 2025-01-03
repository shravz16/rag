import boto3
import json
import os
from langchain_openai import ChatOpenAI
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain import hub
from pinecone import Pinecone
from pinecone import ServerlessSpec
from langchain_pinecone import PineconeVectorStore
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from docx import Document as DocxDocument
import time
import PyPDF2
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain_pinecone import PineconeEmbeddings

class RAGProcessor:
    def __init__(self, queue_url, pinecone_api_key, openai_api_key, region='us-east-2'):
        self.sqs = boto3.client('sqs', region_name=region)
        self.s3 = boto3.client('s3')
        self.queue_url = queue_url
        self.pinecone_api_key = pinecone_api_key
        self.openai_api_key = openai_api_key
        self.setup_pinecone()
        self.setup_embeddings()
        self.setup_llm()
        
    def setup_pinecone(self):
        self.pc = Pinecone(api_key=self.pinecone_api_key)
        self.index_name = "docs-rag-chatbot"
        indexes = self.pc.list_indexes()
        if self.index_name not in [idx['name'] for idx in indexes]:
            self.pc.create_index(
                name=self.index_name,
                dimension=1536,
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1")
            )
        self.index = self.pc.Index(self.index_name)
        
    def setup_embeddings(self):
         self.embeddings = OpenAIEmbeddings(
        model="text-embedding-ada-002",
        openai_api_key=self.openai_api_key)
        
    def setup_llm(self):
        retrieval_qa_chat_prompt = hub.pull("langchain-ai/retrieval-qa-chat")
        self.llm = ChatOpenAI(
            openai_api_key=self.openai_api_key,
            model_name='gpt-4',
            temperature=0.3,  # Reduced for more focused responses
        max_tokens=1000
        )
        
    def extract_text_from_pdf(self, file_path):
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            print(f"Error extracting PDF text: {str(e)}")
            return None
            
    def extract_text_from_docx(self, file_path):
        doc = DocxDocument(file_path)
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        
    def download_from_s3(self, s3_path):
        try:
            bucket = "rag-docs-31331"
            key = s3_path
            print(f"Downloading from S3: s3://{bucket}/{key}")
            local_path = f"/tmp/{os.path.basename(key)}"
            self.s3.download_file(bucket, key, local_path)
            return local_path
        except Exception as e:
            print(f"S3 error: bucket={bucket}, key={key}, error={str(e)}")
            # List bucket contents to debug
            response = self.s3.list_objects_v2(Bucket=bucket)
            print("Available files in bucket:")
            for obj in response.get('Contents', []):
                print(f"- {obj['Key']}")
            raise
        
    def process_document(self, document_text, source):
        cleaned_text = document_text.replace('\n● ', '. ').replace('\n', ' ').strip()
        text_splitter = RecursiveCharacterTextSplitter(
          chunk_size=2000,  # Increased from 500
    chunk_overlap=500,
            separators=[". ", "●", "\n", " ", ""]
        )
        
        doc_chunks = text_splitter.split_text(document_text)
        
        documents = [Document(
    page_content=chunk, 
    metadata={
        "source": source,
        "text": chunk  # Add text content
    }
) for chunk in doc_chunks]
        print(documents)
        
        docsearch = PineconeVectorStore.from_documents(
            documents=documents,
            index_name=self.index_name,
            embedding=self.embeddings,
            namespace="plain-docx-namespace"
        )
        print('passed till doc inn pd 3')
        retriever = docsearch.as_retriever(  search_kwargs={
        "k": 4 
    })
        combine_docs_chain = create_stuff_documents_chain(self.llm, prompt=hub.pull("langchain-ai/retrieval-qa-chat"))
        self.retrieval_chain = create_retrieval_chain(retriever, combine_docs_chain)
        
    def answer_query(self, query):
        response = self.retrieval_chain.invoke({"input": query})
        return response['answer']
        
    def process_messages(self):
        response = self.sqs.receive_message(
            QueueUrl=self.queue_url,
            MaxNumberOfMessages=1,
            WaitTimeSeconds=20
        )
        
        if 'Messages' in response:
            for message in response['Messages']:
                try:
                    body = json.loads(message['Body'])
                    
                    doc_location = body['documentLocation']
                    
                    query = body.get('query', 'where did shravan work?')
                    
                    local_file = self.download_from_s3(doc_location)
                    print(local_file)
                    if local_file.endswith('.pdf'):
                        document_text = self.extract_text_from_pdf(local_file)
                    elif local_file.endswith('.docx'):
                        document_text = self.extract_text_from_docx(local_file)
                    else:
                        raise ValueError("Unsupported file type")
                    print('passed till this')
                    self.process_document(document_text, local_file)
                    query = ['where did shravan work?', 'does shravan know cloud?','what are his projects?','did shravan work in threads?','does shravan have 2 years of work experience?']
                    for q in query:
                        answer = self.answer_query(q)
                        print(f"Query: {q}")
                        print(f"Answer: {answer}")
                    
                    self.sqs.delete_message(
                        QueueUrl=self.queue_url,
                        ReceiptHandle=message['ReceiptHandle']
                    )
                    
                    os.remove(local_file)
                    return True
                except Exception as e:
                    print(f"Error processing message: {str(e)}")
                    return False
        return False

# Initialize and run
processor = RAGProcessor(
    queue_url="https://sqs.us-east-2.amazonaws.com/339712897205/rag",
    pinecone_api_key="1f784203-27de-487e-90be-b2bff7b3b8f8",
    openai_api_key=os.environ.get('OPENAI_API_KEY')
)

while True:
    if not processor.process_messages():
        time.sleep(5)
